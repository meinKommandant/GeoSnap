import openpyxl
import math
import os
import io
from openpyxl.styles import Font, Border, Side
import simplekml
from pathlib import Path
from PIL import Image, ImageOps
import pillow_heif
import shutil

# python-docx imports for WordReportGenerator
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Register HEIF opener
pillow_heif.register_heif_opener()

from .constants import (
    EXCEL_HEADERS,
    COLUMN_WIDTHS,
    KML_CAMERA_ICON,
    ARROW_COLOR,
    ARROW_WIDTH,
    ARROW_MAIN_AXIS_LENGTH,
    ARROW_WING_LENGTH,
    ARROW_WING_ANGLE,
)


class ExcelReportGenerator:
    def __init__(self, title="Listado de Fotos"):
        self.wb = openpyxl.Workbook()
        self.ws = self.wb.active
        self.ws.title = title
        self.thin_border = Border(
            left=Side(style="thin"), right=Side(style="thin"), top=Side(style="thin"), bottom=Side(style="thin")
        )
        self._setup_headers()

    def _setup_headers(self):
        headers = EXCEL_HEADERS

        for cell_coord, text in headers.items():
            cell = self.ws[cell_coord]
            cell.value = text
            cell.font = Font(bold=True)
            cell.border = self.thin_border

        # Column widths
        dims = COLUMN_WIDTHS
        for col, width in dims.items():
            self.ws.column_dimensions[col].width = width

    def add_row(self, row_idx, numero_orden, metadata, altitude_val):
        # Check for No-GPS (dummy coordinates 0.0, 0.0)
        lat = metadata.coordinates.latitude
        lon = metadata.coordinates.longitude
        alt = altitude_val

        if lat == 0.0 and lon == 0.0:
            lat = ""
            lon = ""
            alt = ""

        cells = [
            (2, numero_orden),
            (3, metadata.filename),
            (4, ""),
            (5, str(metadata.timestamp)),
            (6, lat),
            (7, lon),
            (8, alt),
            (9, metadata.coordinates.azimuth if metadata.coordinates.azimuth is not None else ""),
        ]
        for col_idx, val in cells:
            c = self.ws.cell(row=row_idx, column=col_idx, value=val)
            c.border = self.thin_border

    def save(self, path):
        self.wb.save(str(path))


class KmzReportGenerator:
    def __init__(self, thumbs_dir: Path):
        self.kml = simplekml.Kml()
        self.thumbs_dir = thumbs_dir
        self._prepare_thumbs_dir()

    def _prepare_thumbs_dir(self):
        if self.thumbs_dir.exists():
            shutil.rmtree(self.thumbs_dir)
        self.thumbs_dir.mkdir(parents=True, exist_ok=True)

    def add_point(self, numero_orden, metadata, img_path, altitude_val):
        # Priority: Excel Sequence ID > Fallback to Loop Counter
        display_id = metadata.sequence_id if metadata.sequence_id else numero_orden
        titulo_punto = f"Foto Nº {display_id}"

        # Arrow Logic (Azimuth)
        if metadata.coordinates.azimuth is not None:
            # Create Placemark with MultiGeometry
            pnt = self.kml.newmultigeometry(name=titulo_punto)

            # 1. Original Point (Camera)
            # When calling newpoint on a MultiGeometry, it adds the point to the collection
            pnt.newpoint(coords=[(metadata.coordinates.longitude, metadata.coordinates.latitude)])

            # 2. Flecha Amarilla
            lat = metadata.coordinates.latitude
            lon = metadata.coordinates.longitude
            az = metadata.coordinates.azimuth

            # Calcular puntos
            end_lat, end_lon = self._calculate_dest_point(lat, lon, ARROW_MAIN_AXIS_LENGTH, az)
            w1_lat, w1_lon = self._calculate_dest_point(end_lat, end_lon, ARROW_WING_LENGTH, az + ARROW_WING_ANGLE)
            w2_lat, w2_lon = self._calculate_dest_point(end_lat, end_lon, ARROW_WING_LENGTH, az - ARROW_WING_ANGLE)

            # Draw line: Start -> End -> Wing1 -> End -> Wing2
            arrow_coords = [(lon, lat), (end_lon, end_lat), (w1_lon, w1_lat), (end_lon, end_lat), (w2_lon, w2_lat)]

            ls = pnt.newlinestring(coords=arrow_coords)
            ls.style.linestyle.color = ARROW_COLOR
            ls.style.linestyle.width = ARROW_WIDTH
        else:
            # Crear Placemark con Point
            pnt = self.kml.newpoint(name=titulo_punto)
            pnt.coords = [(metadata.coordinates.longitude, metadata.coordinates.latitude)]

        # Icon style (Red camera)
        # http://maps.google.com/mapfiles/kml/pal4/icon46.png is a camera
        pnt.style.iconstyle.icon.href = KML_CAMERA_ICON
        pnt.style.iconstyle.color = simplekml.Color.red  # Red tint

        # Miniatura
        thumb_filename = f"thumb_{img_path.name}"
        local_thumb_path = self.thumbs_dir / thumb_filename
        img_src_in_kmz = ""

        try:
            with Image.open(img_path) as img:
                img = ImageOps.exif_transpose(img)
                img.thumbnail((800, 800))
                img.save(local_thumb_path, quality=75)
            added_path = self.kml.addfile(str(local_thumb_path))
            if added_path:
                img_src_in_kmz = added_path.replace("\\", "/")
        except Exception:
            pass

        # HTML Description (Data table)
        # Data: Nº, File, Description, Date, Latitude, Longitude, Altitude
        desc_text = (metadata.description or "").strip() if hasattr(metadata, "description") else ""
        table_html = f"""
        <table border="1" style="border-collapse: collapse; width: 100%;">
            <tr><td><b>Nº</b></td><td>{display_id}</td></tr>
            <tr><td><b>Archivo</b></td><td>{metadata.filename}</td></tr>
            <tr><td><b>DESCRIPCIÓN</b></td><td>{desc_text}</td></tr>
            <tr><td><b>Fecha</b></td><td>{metadata.timestamp}</td></tr>
            <tr><td><b>Latitud</b></td><td>{metadata.coordinates.latitude}</td></tr>
            <tr><td><b>Longitud</b></td><td>{metadata.coordinates.longitude}</td></tr>
            <tr><td><b>Altitud [m]</b></td><td>{altitude_val}</td></tr>
            <tr><td><b>Rumbo [°]</b></td><td>{metadata.coordinates.azimuth if metadata.coordinates.azimuth is not None else ""}</td></tr>
        </table>
        """

        img_html = (
            f'<img src="{img_src_in_kmz}" style="max-width:400px; display:block; margin-bottom:10px;"/>'
            if img_src_in_kmz
            else ""
        )

        # In the balloon we show only the photo and the table
        pnt.description = f"{img_html}{table_html}"

    def save(self, path):
        self.kml.savekmz(str(path))

    def cleanup(self):
        if self.thumbs_dir.exists():
            shutil.rmtree(self.thumbs_dir)

    def _calculate_dest_point(self, lat, lon, dist_m, bearing_deg):
        """Calculates destination point given origin point, distance (m) and bearing (degrees)."""
        R = 6378137  # Earth radius in meters
        brng = math.radians(bearing_deg)
        lat1 = math.radians(lat)
        lon1 = math.radians(lon)

        lat2 = math.asin(math.sin(lat1) * math.cos(dist_m / R) + math.cos(lat1) * math.sin(dist_m / R) * math.cos(brng))
        lon2 = lon1 + math.atan2(
            math.sin(brng) * math.sin(dist_m / R) * math.cos(lat1),
            math.cos(dist_m / R) - math.sin(lat1) * math.sin(lat2),
        )

        return math.degrees(lat2), math.degrees(lon2)


class WordReportGenerator:
    """
    Generates a Word document report with photos in a 2-column journalistic layout.
    Landscape orientation with 4 photos per page (2 per column).
    Includes EXIF rotation fix and centered/bold captions.
    """

    def __init__(self, title="Reporte Fotográfico"):
        self.doc = Document()

        # --- Configure Section: Landscape and 2 Columns ---
        section = self.doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)

        # Margins (1.5 cm to maximize vertical space)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

        # Force 2 columns
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0] if sectPr.xpath('./w:cols') else OxmlElement('w:cols')
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '720')  # ~1.27cm
        if not sectPr.xpath('./w:cols'):
            sectPr.append(cols)

    def add_photo(self, numero_orden, metadata, img_path):
        """
        Adds photo (with EXIF rotation fix) and centered/black caption.
        """
        display_id = metadata.sequence_id if metadata.sequence_id else str(numero_orden)
        desc_text = metadata.description.strip() if metadata.description else "XXXXXXXXXXXXX"

        # 1. Insert Image (with EXIF orientation fix)
        p_img = self.doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.keep_with_next = True
        p_img.paragraph_format.space_before = Pt(2)
        p_img.paragraph_format.space_after = Pt(2)

        run_img = p_img.add_run()

        if img_path and os.path.exists(img_path):
            try:
                # Process image with Pillow to rotate according to EXIF
                with Image.open(img_path) as img:
                    img_fixed = ImageOps.exif_transpose(img)

                    # Save to memory stream for python-docx
                    img_stream = io.BytesIO()
                    img_fixed.save(img_stream, format='JPEG', quality=85)
                    img_stream.seek(0)

                    # Width 9.5 cm for landscape format (2 cols)
                    run_img.add_picture(img_stream, width=Cm(9.5))
            except Exception as e:
                self._add_placeholder(run_img, f"[ERROR IMAGEN: {str(e)}]")
        else:
            self._add_placeholder(run_img, "[IMAGEN NO ENCONTRADA]")

        # 2. Insert Caption (Centered and Black)
        p_cap = self.doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)

        # Part 1: "Figura X.- " in Bold
        run_bold = p_cap.add_run(f"Figura {display_id}.- ")
        run_bold.font.name = "Calibri"
        run_bold.font.size = Pt(10)
        run_bold.font.bold = True
        run_bold.font.color.rgb = RGBColor(0, 0, 0)  # Black

        # Part 2: Description normal
        run_desc = p_cap.add_run(desc_text)
        run_desc.font.name = "Calibri"
        run_desc.font.size = Pt(10)
        run_desc.font.bold = False
        run_desc.font.color.rgb = RGBColor(0, 0, 0)  # Black

        # Red warning only for placeholder text
        if desc_text == "XXXXXXXXXXXXX":
            run_desc.font.color.rgb = RGBColor(255, 0, 0)

    def _add_placeholder(self, run, text):
        """Draws visible text when image is missing."""
        run.add_text(f"\n{text}\n")
        run.font.bold = True
        run.font.color.rgb = RGBColor(200, 0, 0)  # Dark Red
        run.font.size = Pt(12)

    def save(self, path):
        self.doc.save(str(path))
